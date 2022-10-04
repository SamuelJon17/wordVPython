import os
import streamlit as st
from docx import Document
import re

import pandas as pd
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from titlecase import titlecase

# Reference
# https://docs.streamlit.io/library/api-reference/session-state
# https://discuss.streamlit.io/t/how-to-get-multiple-inputs-using-same-text-input-box/28435/2 

#=========================================================================
#                       Webpage Header
#=========================================================================
st.title('Word Document V Python')
st.subheader('Upload Word Document(s) for Processing')
files = st.file_uploader('Multiple Allowed', accept_multiple_files=True)
path = st.text_input("Please Enter Path to Save", placeholder = "/User/Downloads/", key = "path")

if "findReplace" not in st.session_state:
    st.session_state.findReplace = []

#=========================================================================
#                       findReplace Methods
#=========================================================================

def findReplaceExact(String, find, replace, caseSens = False):
    """
    Find and replace text in String
    @param String: String from either a Paragraph or Runner object
    @param find: Substring wanting to replace
    @param replace: Replacement string
    @param caseSens: True - if case sensitive, False otherwise
    @return String, original string if find not in String otherwise updated String with replace inserted
    """
    if not caseSens:
        regex = r"(?i)(\b" + str(find) + r"\b)"
        if len(re.findall(regex, String))>0:
            result = re.split(regex, String)
            return ''.join([replace if i.lower() == find.lower() else i for i in result])
        else:
            return String
    else:
        regex = r"(\b" + str(find) + r"\b)"
        if len(re.findall(regex,String))>0:
            result = re.split(regex, String)
            return ''.join([replace if i.lower() == find.lower() else i for i in result])
        else:
            return String

def findReplaceExactOverDocument(document, find, replace, caseSens = False):
    """
    Find and replace text in Entire document
    @param document: Document object
    @param find: Substring wanting to replace
    @param replace: Replacement string
    @param caseSens: True - if case sensitive, False otherwise
    @return String, original string if find not in String otherwise updated String with replace inserted
    """
    for paragraph in document.paragraphs:
        if findReplaceExact(paragraph.text, find, replace, caseSens=caseSens) != paragraph.text:
            st.write("\n=============================")
            st.write("**The following statement:** ")
            st.write(f"\n\t'{paragraph.text}'")
            for run in paragraph.runs:
                if (findReplaceExact(run.text, find, replace, caseSens=caseSens) != run.text):
                    run.text = findReplaceExact(run.text, find, replace, caseSens=caseSens)
            st.write(f"**Was replaced to:**")
            st.write(f"\n\t'{paragraph.text}'")
    pass

def add_callback():
    # Handle if find already in list
    findIndex = [i for i, x in enumerate(st.session_state.findReplace) if x['find'] == st.session_state.find]
    if len(findIndex) >0:
        st.session_state.findReplace[findIndex[0]].update({"find":st.session_state.find, "replace":st.session_state.replace,
        "caseSensitive":st.session_state.case})
    
    # Otherwise update list
    else:
        st.session_state.findReplace.append({"find":st.session_state.find, "replace":st.session_state.replace,
        "caseSensitive":st.session_state.case})
    

def delete_callback():
    # Remove last added index
    if len(st.session_state.findReplace)==0:
        pass
    else:
        st.session_state.findReplace = st.session_state.findReplace[:-1]

def saveDocs(documents, docNames, extension = ""):
    if len(documents) ==0:
        st.subheader("Please upload a document")
    else:
        for i in range(len(documents)):
            docNames[i] = docNames[i].split(".docx")[0] + "_" + extension + "Processed.docx"
            documents[i].save(os.path.join(path,docNames[i]))
            st.write(f"**File successfully saved at: {os.path.join(path,docNames[i])}**")
        st.subheader("Please refresh page to start a new session".title())
    pass

#=========================================================================
#                       Stylizer Methods
#=========================================================================
def findCaption(String):
    """
    Finding captions titled Figure(s) #: or Table(s) #: 
    @param obj: String, generally the Paragraph 's object text (Paragraph.text)
    @return String, found caption
    """
    found = re.search(r"(?i)(figure\s+\d+:?)", String) ## case insensitive
    if found != None:
        return found.group(1).strip(":")

    found = re.search(r"(?i)(table\s+\d+:?)", String) ## case insensitive
    if found != None:
        return found.group(1).strip(":")
    return None

def changeTextStyle(obj, bold = False, alignText = False, alignTable = False, fontName = "Times New Roman", size = 12, p = False):
    """
    Main funciton to change text style (bold, alignment, font, size)
    @param obj: Paragraph or Runner object
    @returns: void, changes style in the object
    """
    if p:
        return "The following document updated the font to " + fontName +", " +str(size)+ " size, bolded where appropriate," +\
            "and aligned Justify"
    
    try:
        obj.font.name = fontName ## Times New Roman
        obj.font.size = docx.shared.Pt(size) ## font 12
        obj.font.bold = bold

    except:
        obj.style.font.name = fontName ## Times New Roman
        obj.style.font.size = docx.shared.Pt(size) ## font 12
        obj.bold = bold
        
    if alignText:
        obj.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY ## justifed alignment
    if alignTable:
        obj.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY ## change if requiring something else, can only align whole table,
                                                   ## not just certain rows/columns/cells      
    
    pass

def changeStyleMain(obj):
    """
    Main function to change style of each Style_id (headings, captions, all other)
    This assumes the end-user is appropriately using Word's style formatter
    @param obj: Paragraph object
    @returns: usedCaptions and usedCaptionsInText, changes style in the object(s)
    """
    
    usedCaptions = set()
    usedCaptionsInText = set()
    
    ## Headers (heading 1, 2, 3, ...)
        # These can be more customaizable (i.e. heading 1 style != heading 2 style ...)
        # Heading 2 in the reference document had its own styling in comparison to Heading 1 & 3 which is why
        # it is skipped in this seciton
    if ("heading" in obj.style.style_id.lower() and obj.style.style_id != "Heading2"): 
        for i in obj.runs:
            changeTextStyle(i, bold = True, alignText = False)
        obj.text = titlecase(obj.text)
    
    ## Heading 2 specific, different font size and boldness
    elif (obj.style.style_id == "Heading2"): 
        for i in obj.runs:
            changeTextStyle(i, bold = False, size = 10, alignText = False)
        obj.text = titlecase(obj.text)
    
    ## Captions (Figures/Tables)
        # Assumes end-user has "Figure Caption" or "Caption" style for captions, otherwise will not work
    elif (obj.style.style_id == "Caption" or obj.style.style_id=="FigureCaption"):   
        # Find captions in Captions style
        foundCaption = findCaption(obj.text)
        end = ""
        if foundCaption != None:
            usedCaptions.add(foundCaption.title())
            if ":" not in obj.text.split(foundCaption,1)[1]:
                st.write(f"The following caption does not follow style guidelines: {obj.text}")
        
        # For captions, bold either Figure or Table, its respective number, and unbold the semicolon
        changeTextStyle(obj, bold = False, alignText = False) ## 
        for i in obj.runs:
            if ("figure" in i.text.lower() or "table" in i.text.lower()):
                ## Logically, this works and is updatin styles but output is not reflected for some reason
                changeTextStyle(i, bold = True, alignText = False)
            else:
                ## Logically, this works and is updatin styles but output is not reflected for some reason
                changeTextStyle(i, bold = False, alignText = False)
        
        if (obj.text != titlecase(obj.text)):
            st.write(f"The following caption is not properly capitalized: {obj.text}")

    
    ## All other body of text
        # Style != Heading 1, 2, 3, ... & Caption
        # Includes list paragraphs, body text, normal, etc. styles
    else:
        # Find captions in Style != Caption
        foundCaption = findCaption(obj.text)
        if foundCaption != None:
            ## If we care about references made prior to Caption use below otherwise refer to next step
            #if foundCaption not in usedCaptions:
            #    print(foundCaption + " was referenced before caption")
            
            ## Keep track of what's referenced in text (Style != Caption)
            usedCaptionsInText.add(foundCaption.title())
        
        # For text, change style
        for i in obj.runs:
            if ("note:" in i.text.lower() or "notes:" in i.text.lower()):
                changeTextStyle(i, bold = True, alignText = True)
                st.write("'Note:' in the following string was bolded: {}".format(obj.text))
            elif (i.font.color.type ==1):
                continue
            else:
                changeTextStyle(i, bold = False, alignText = True)
    return usedCaptions, usedCaptionsInText

def checkTableOrdering(table):
    """
    Function to check if the table ordering is correct for experiment numbers, parameters, acronyms, etc.
    Will return a dictionary with the table number and whether it passed ordering or not
    @param table - Table object
    @return - dictionary <String,String>, stating if ordering for experiment numbers or parameters, 
                acronyms, etc. are correct
    
    """
    changeTextStyle(table, alignTable = True) ## update style of entire table
                                              ## Cannot change style of row, column or cell individually
    
    ## Create DataFrame
    df = [['' for i in range(len(table.columns))] for j in range(len(table.rows))]
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):                
            if cell.text:
                df[i][j] = cell.text
    df = pd.DataFrame(df)
    df = df.rename(columns=df.iloc[0].str.strip()).drop(df.index[0])
    
    ## Focus on DataFrame's that have these column headers in acceptList
        # If column header found, check if that column is appropriately sorted
    acceptList = ["Experiment", "Experiments", "Parameter","Parameters","Acronym","Acronyms","Experiment Number"]
    acceptValue = [a for a in df.columns if a in acceptList]
    if len(acceptValue) >0:
        if list(df[acceptValue[0]]) == sorted(list(df[acceptValue[0]])):
            return "pass"
        else:
            return "fail"
    return "n.a."


def main(document):
    """
    @param document: Document object
    @return, void, make changes for document text and tables, print if ordering is correct for tables and if figures
            and tables are referenced at least once in body
    """
    
    ## Text
    usedCaptions = set()
    usedCaptionsInText = set()
    for k in document.paragraphs:
        a,b = changeStyleMain(k)
        usedCaptions.update(a)
        usedCaptionsInText.update(b)

    unusedCaptions = usedCaptions.symmetric_difference(usedCaptionsInText)
    for i in unusedCaptions:
        st.write(i + " was never referenced in the document.")
    
    ## Tables
    checkTableDict = {}
    count = 1
    for table in document.tables:
        checkTableDict.update({"table" + str(count) : checkTableOrdering(table)})
        count+= 1 
    for key, value in checkTableDict.items():
        if value == "fail":
            st.write(key+" is not ordered properly")
    
    st.write(changeTextStyle(1, p=True))
    st.write(f"**Table of Content cross referencing needs to be updated.**")
    st.write(f"**Cross referencing for non-captions may have broken in the process**")
    
#=========================================================================
#                       Webpage Contents
#=========================================================================

tab1, tab2 = st.tabs(["findReplace", "stylize"])
with tab1:
    ## Update state session find/replace statements
    with st.form(key='myForm'):
        st.subheader("Find And Replace Batch Process")
        find = st.text_input("Find", placeholder="Add find word", key = "find")
        replace = st.text_input("Replace", placeholder = "Add replace word", key = "replace")
        caseSensitive = st.checkbox("Case Sensitive", value = False, key = "case")
        submit_button = st.form_submit_button(label='Add', on_click=add_callback)
        delete_button = st.form_submit_button(label="Remove Last", on_click=delete_callback)

    ## Print state session find/replace statements
    count=1
    for obj in st.session_state.findReplace:
        st.write(f"{count}. Find: {obj['find']}  | Replace: {obj['replace']}  | Case Sesnsitive: {obj['caseSensitive']}\n")
        count+=1


    ## Now: This will loop through each document and within each document loop again for each findReplace item
    ## Future: This can be improved by looping through document only one pass through and check all findReplace at once
    ## Future:  Save JSON of "Find/Replace/CaseSens" to import rather than typing it in each time
    if st.button("Process Documents for Find/Replace"):
            if files is not None:
                docNames=[]
                documents = []
                for document in files:
                    docName = str(document.name)
                    docNames.append(docName)
                    st.subheader(f"{docName}")
                    document = Document(document)
                    documents.append(document)
                    for item in st.session_state.findReplace:
                        findReplaceExactOverDocument(document, item['find'], item['replace'], item['caseSensitive'])
                saveDocs(documents, docNames, "findReplace")


with tab2:
    st.subheader("Document Style Guidelines")
    if st.button("Process Documents for Stylize"):
            if files is not None:
                docNames=[]
                documents = []
                for document in files:
                    docName = str(document.name)
                    docNames.append(docName)
                    st.subheader(f"{docName}")
                    document = Document(document)
                    documents.append(document)
                    main(document)
                saveDocs(documents, docNames, "Style")